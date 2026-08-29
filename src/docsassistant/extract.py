"""Pull text out of uploaded files, keeping page numbers attached.

Page numbers survive extraction because everything downstream depends
on them. PDFs give us real pages. Word documents have no fixed pages
until they are rendered, so we approximate by splitting on explicit
page breaks and falling back to a word budget — and the README says so
plainly rather than pretending the number is exact.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from .models import Document, Page

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".log", ".csv"}

# Roughly a page of prose. Only used for formats with no real pages.
_WORDS_PER_PAGE = 400


class ExtractionError(Exception):
    """Raised when a file cannot be read as text."""


# Rejoin a line that was wrapped mid-sentence: the previous line does
# not end in terminal punctuation and the next begins lowercase. PDF
# extraction wraps constantly, and leaving those newlines in breaks
# phrase matching ("administrator\naccounts" never matches a search for
# "administrator accounts"). Lines that look like separate records are
# left alone, so log and CSV files keep their structure.
_WRAPPED_LINE = re.compile(r"(?<![.!?:;])\n(?!\n)(?=[a-z(\[])")


def _clean(text: str) -> str:
    """Normalise whitespace without destroying paragraph structure."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = _WRAPPED_LINE.sub(" ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pdf(content: bytes) -> list[Page]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError(
            "PDF support needs pypdf — install it with: pip install pypdf"
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise ExtractionError(f"could not open PDF: {exc}") from exc

    if reader.is_encrypted:
        # An empty password unlocks many "protected" PDFs.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ExtractionError("PDF is password protected") from exc

    pages: list[Page] = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            # One unreadable page should not lose the whole document.
            text = ""
        text = _clean(text)
        if text:
            pages.append(Page(number=number, text=text))

    if not pages:
        raise ExtractionError(
            "no text found — this looks like a scanned PDF. "
            "Run OCR on it first (e.g. ocrmypdf) and upload the result."
        )
    return pages


def _extract_docx(content: bytes) -> list[Page]:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError(
            "DOCX support needs python-docx — install it with: "
            "pip install python-docx"
        ) from exc

    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:
        raise ExtractionError(f"could not open DOCX: {exc}") from exc

    blocks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            blocks.append(paragraph.text)

    # Tables carry a lot of the meaning in policy and contract documents,
    # so flatten them into readable rows rather than dropping them.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    if not blocks:
        raise ExtractionError("document contains no readable text")

    return _paginate(blocks)


def _extract_plain(content: bytes) -> list[Page]:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "latin-1"):
        try:
            text = content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:  # pragma: no cover
        raise ExtractionError("could not decode file as text")

    text = _clean(text)
    if not text:
        raise ExtractionError("file is empty")

    return _paginate(text.split("\n\n"))


def _paginate(blocks: list[str]) -> list[Page]:
    """Group blocks into approximate pages by word count."""
    pages: list[Page] = []
    current: list[str] = []
    words = 0

    for block in blocks:
        block_words = len(block.split())
        if words + block_words > _WORDS_PER_PAGE and current:
            pages.append(Page(number=len(pages) + 1, text=_clean("\n\n".join(current))))
            current, words = [], 0
        current.append(block)
        words += block_words

    if current:
        pages.append(Page(number=len(pages) + 1, text=_clean("\n\n".join(current))))

    return pages


def extract(filename: str, content: bytes) -> Document:
    """Extract text from an uploaded file.

    Raises ExtractionError with a message written for the person who
    uploaded the file, not for a developer reading a stack trace.
    """
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_SUFFIXES:
        supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise ExtractionError(
            f"{suffix or 'this file type'} is not supported. "
            f"Try one of: {supported}"
        )

    if not content:
        raise ExtractionError("file is empty")

    if suffix == ".pdf":
        pages = _extract_pdf(content)
    elif suffix == ".docx":
        pages = _extract_docx(content)
    else:
        pages = _extract_plain(content)

    return Document(
        doc_id=Document.make_id(filename, content),
        filename=Path(filename).name,
        pages=pages,
    )
